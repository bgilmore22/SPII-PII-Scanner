import re #Imports the regular expression module so that we can use it for pattern searching in the files
import os #We import the os module so that we can manage the different file extensions
from docx import Document #This is imporoted so that we can read and edit word's docx documents
from PyPDF2 import PdfReader #This is imported so we can read pdf files
from fpdf import FPDF #This is the import that allows us to edit the pdf files

#These are the patterns that we are looking for in the files that contain sensitive information
sensitiveInfo = {
    "Email": r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", #Email addresses with the format laid out for how a normal email address is
    "Credit Card": r"\b(?:\d[ -]*?){13,16}\b", #This is the standardized credit card format
    "TFN": r"\b\d{3} \d{3} \d{3}\b", #A pattern search for the Australian resident's tax file number
    "Phone": r"\b04\d{8}\b", #Format to search for phone numbers
    "ABN": r"\b\d{2} \d{3} \d{3} \d{3}\b", #The number that a business uses to identify themselves
}

#This is the helper function that will look for the previously defined patterns in the files and then replace any instance matched with [REDACTED]
def redact_text(text):
    #This is an empty dictionary that contains all the patterns found in the files
    findings = {label: re.findall(pattern, text) for label, pattern in sensitiveInfo.items()}
    #We then copy the text to a different variable that we can use to apply the redaction on
    redacted_text = text
    #We then loop through each item in the findings and match them to the patterns
    for label, matches in findings.items():
        #If a match has been found, we then edit the file to replace the information with [REDACTED]
        for match in matches:
            redacted_text = redacted_text.replace(match, "[REDACTED]")
    #We then return both the redacted text and the dictionary
    return redacted_text, findings

#This is a helper function that is used when the inputted file has a file extension of docx
def process_docx(path):
    #We first open and read the file
    doc = Document(path)
    #We then combine all the information, characters, numbers etc in the file into one long string
    full_text = "\n".join([p.text for p in doc.paragraphs])
    #We then call the redacting helper function on the singular string
    redacted_text, findings = redact_text(full_text)
    #A new document is then created with the redacted content
    redacted_doc = Document()
    for line in redacted_text.split("\n"):
        redacted_doc.add_paragraph(line)
    #The newly created redacted document is also renamed to include _redacted so that the user knows it has been processed
    redacted_path = path.replace(".docx", "_redacted.docx")
    redacted_doc.save(redacted_path)
    return findings, redacted_path

#This helper function is specific to processing PDF files
def process_pdf(path):
    #We first open and read the PDF file
    reader = PdfReader(path)
    text = ""
    #We then run a loop through the file so that we are able to extract all the information from each page
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    #Once the information is extracted, we call the redaction helper function
    redacted_text, findings = redact_text(text)
    #We then create a new PDF document to save the redacted output to
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    #This is the loop that adds the redacted text to the redacted file
    for line in redacted_text.split('\n'):
        pdf.multi_cell(0, 10, line)
    #We then save the newly redacted file with the _redacted added to the name so that the user knows it has been processed
    redacted_path = path.replace(".pdf", "_redacted.pdf")
    pdf.output(redacted_path)
    return findings, redacted_path

#This is the main function that will scan a file and redact sensitive information
def scan_file(input_path):
    #We first must find out what file extension the file has
    ext = os.path.splitext(input_path)[-1].lower()
    #If the file has a docx extension, we call the helper function specifically made for this extension
    if ext == ".docx":
        findings, redacted_path = process_docx(input_path)
    #If the extension is pdf, we call the pdf helper function
    elif ext == ".pdf":
        findings, redacted_path = process_pdf(input_path)
    #If docx or pdf isnt given, then the program cant support it and read its information and redact it
    else:
        raise ValueError("Unsupported file type. Use .pdf or .docx")
    
    #We then print the summary of the information we redacted
    print("\nSummary:")
    risk_score = 0
    summary = {}

    #This loop is used to keep track of how many redaction occurences occured so we can assign a risk rating to the file
    for label, matches in findings.items():
        count = len(matches)
        summary[label] = count
        if count > 0:
            #A higher count is given to credit card information or abn information as they are extremely dangerous
            if label in ["Credit Card", "ABN"]:
                risk_score += count * 2
            else:
                risk_score += count
        #This is then the output format for the risk ratings
        print(f"{label}: {count} found")

    #We then run the total score through a if else if loop to determine its total risk level
    if risk_score == 0:
        overall_risk = "None"
    elif risk_score <= 2:
        overall_risk = "Low"
    elif risk_score <= 5:
        overall_risk = "Medium"
    else:
        overall_risk = "High"

    #This tells the user the final decided risk rating for the file and also the location that the new file was stored in
    print(f"\n🔒 Overall File Risk Rating: {overall_risk.upper()} (Score: {risk_score})")
    print(f"\nRedacted file saved as:\n{redacted_path}")

    #Finally we return a dictionary containing a summary of what was found
    return {
        "summary": summary,
        "risk_score": risk_score,
        "risk_level": overall_risk,
        "output_path": redacted_path
    }
